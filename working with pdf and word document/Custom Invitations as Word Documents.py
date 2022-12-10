import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
guest = ['Prof. Plum','Miss Scarlet','Col. Mustard','Al Sweigart','RoboCop']
doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Brush Script MT'
font.size = Pt(20)
for i in guest:
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.paragraphs[len(doc.paragraphs)-1].runs[0].italic = True
    doc.add_paragraph(i)
    doc.paragraphs[len(doc.paragraphs)-1].runs[0].bold = True
    doc.paragraphs[len(doc.paragraphs) - 1].runs[0].font.name = 'Times New Roman'
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    doc.paragraphs[len(doc.paragraphs)-1].runs[0].italic = True
    doc.add_paragraph('April 1st')
    doc.paragraphs[len(doc.paragraphs) - 1].runs[0].font.name = 'Times New Roman'
    doc.add_paragraph("at 7 o'clock")
    doc.paragraphs[len(doc.paragraphs)-1].runs[0].italic = True
    doc.paragraphs[len(doc.paragraphs)-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

for paragraph in doc.paragraphs:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save('main.docx')